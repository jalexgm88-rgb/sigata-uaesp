"""
modules/reportes.py
--------------------------------------------------------------------------------
Generacion automatica de reportes ejecutivos en PDF, Excel y Word a partir del
DataFrame de acciones afirmativas ya filtrado, respetando siempre los filtros
inteligentes aplicados por el usuario en el momento de la exportacion.
--------------------------------------------------------------------------------
"""

import math
import textwrap
from datetime import datetime
from io import BytesIO

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from fpdf import FPDF

from config.settings import (
    APP_ENTITY, APP_FULL_NAME, APP_NAME, APP_VERSION, LOGO_PATH,
    PROYECTO_INVERSION, USUARIO_CONECTADO,
)
from modules import analisis
from modules.ui import fecha_larga_es


# Equivalencias en texto plano para simbolos Unicode usados en la interfaz
# (flechas de tendencia, semaforos) que las fuentes core de PDF (Helvetica)
# no pueden representar, ya que solo soportan el conjunto de caracteres
# Latin-1/cp1252.
_PDF_EQUIVALENCIAS_UNICODE = {
    "▲": "+", "▼": "-", "▬": "=",
    "🟢": "", "🟡": "", "🔴": "",
    "–": "-", "—": "-", "’": "'", "“": '"', "”": '"',
}


def _pdf_safe_text(texto: str) -> str:
    """
    Convierte un texto arbitrario (que puede provenir de datos cargados por
    el usuario o de textos generados dinamicamente) a una version segura
    para las fuentes core de fpdf2. Sustituye simbolos conocidos por su
    equivalente en texto plano y, como salvaguarda final, reemplaza
    cualquier caracter restante fuera de Latin-1 para evitar que
    `FPDFUnicodeEncodingException` interrumpa la generacion del informe.
    """
    texto = str(texto) if texto is not None else ""
    for simbolo, equivalente in _PDF_EQUIVALENCIAS_UNICODE.items():
        texto = texto.replace(simbolo, equivalente)
    return texto.encode("latin-1", errors="replace").decode("latin-1")


def _write_wrapped(pdf: FPDF, texto: str, ancho_caracteres: int = 100, alto_linea: int = 6):
    """
    Escribe texto envolviendo lineas manualmente con textwrap y dibujando cada
    linea con pdf.cell(). Evita un problema conocido de fpdf2 en el que
    multi_cell() puede lanzar FPDFException ("Not enough horizontal space to
    render a single character") en ciertas combinaciones de texto y ancho.
    El texto se sanea con `_pdf_safe_text()` antes de envolverlo.
    """
    texto = _pdf_safe_text(texto)
    lineas = textwrap.wrap(texto, width=ancho_caracteres) or [""]
    for linea in lineas:
        pdf.cell(0, alto_linea, linea, ln=1)


# --------------------------------------------------------------------------------
# Excel
# --------------------------------------------------------------------------------
def generar_reporte_excel(df: pd.DataFrame, kpis: dict) -> bytes:
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        resumen = pd.DataFrame([
            {"Indicador": "Recicladores registrados", "Valor": kpis.get("total_recicladores", 0)},
            {"Indicador": "Organizaciones", "Valor": kpis.get("total_organizaciones", 0)},
            {"Indicador": "Acciones afirmativas", "Valor": kpis.get("total_acciones", 0)},
            {"Indicador": "Beneficiarios unicos", "Valor": kpis.get("beneficiarios_unicos", 0)},
            {"Indicador": "Cobertura (%)", "Valor": kpis.get("cobertura_pct", 0)},
            {"Indicador": "Presupuesto total", "Valor": kpis.get("presupuesto_total", 0)},
            {"Indicador": "Presupuesto ejecutado", "Valor": kpis.get("presupuesto_ejecutado", 0)},
            {"Indicador": "Presupuesto disponible", "Valor": kpis.get("presupuesto_disponible", 0)},
            {"Indicador": "Acciones ejecutadas este ano", "Valor": kpis.get("acciones_ejecutadas_anio", 0)},
            {"Indicador": "Acciones pendientes", "Valor": kpis.get("acciones_pendientes", 0)},
            {"Indicador": "Alertas activas", "Valor": kpis.get("alertas_activas", 0)},
        ])
        resumen.to_excel(writer, sheet_name="Resumen KPI", index=False)

        columnas_detalle = [
            "fecha", "documento", "nombre", "organizacion", "localidad", "tipo_accion",
            "programa", "proyecto", "responsable", "estado", "presupuesto",
            "presupuesto_ejecutado", "beneficio", "sexo", "edad", "grupo_poblacional",
        ]
        columnas_detalle = [c for c in columnas_detalle if c in df.columns]
        df[columnas_detalle].to_excel(writer, sheet_name="Detalle de acciones", index=False)

        if "localidad" in df.columns and not df.empty:
            por_localidad = df.groupby("localidad").agg(
                acciones=("id", "count"), beneficiarios=("documento", "nunique"),
                presupuesto=("presupuesto", "sum"),
            ).reset_index()
            por_localidad.to_excel(writer, sheet_name="Por localidad", index=False)

    buffer.seek(0)
    return buffer.getvalue()


# --------------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------------
class ReportePDF(FPDF):
    def header(self):
        self.set_fill_color(14, 110, 79)
        self.rect(0, 0, 210, 22, "F")
        self.set_text_color(255, 255, 255)
        self.set_font("Helvetica", "B", 15)
        self.set_xy(10, 6)
        self.cell(0, 8, f"{APP_NAME} - Reporte Ejecutivo", ln=1)
        self.set_font("Helvetica", "", 9)
        self.set_x(10)
        self.cell(0, 6, APP_ENTITY, ln=1)
        self.ln(6)
        self.set_text_color(30, 30, 30)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(120, 120, 120)
        self.cell(0, 10, f"Pagina {self.page_no()} | Generado el {datetime.now().strftime('%Y-%m-%d %H:%M')}", align="C")


def generar_reporte_pdf(df: pd.DataFrame, kpis: dict, recomendaciones: list, analisis_texto: str = "") -> bytes:
    pdf = ReportePDF()
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(14, 110, 79)
    pdf.cell(0, 8, "Indicadores clave", ln=1)
    pdf.set_text_color(30, 30, 30)
    pdf.set_font("Helvetica", "", 10)

    filas_kpi = [
        ("Recicladores registrados", kpis.get("total_recicladores", 0)),
        ("Organizaciones", kpis.get("total_organizaciones", 0)),
        ("Acciones afirmativas", kpis.get("total_acciones", 0)),
        ("Beneficiarios unicos", kpis.get("beneficiarios_unicos", 0)),
        ("Cobertura (%)", f"{kpis.get('cobertura_pct', 0)}%"),
        ("Presupuesto total", f"${kpis.get('presupuesto_total', 0):,.0f}"),
        ("Presupuesto ejecutado", f"${kpis.get('presupuesto_ejecutado', 0):,.0f}"),
        ("Acciones pendientes", kpis.get("acciones_pendientes", 0)),
        ("Alertas activas", kpis.get("alertas_activas", 0)),
    ]
    for etiqueta, valor in filas_kpi:
        pdf.cell(90, 7, etiqueta, border=1)
        pdf.cell(0, 7, str(valor), border=1, ln=1)

    pdf.ln(4)
    if analisis_texto:
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(14, 110, 79)
        pdf.cell(0, 8, "Analisis ejecutivo", ln=1)
        pdf.set_text_color(30, 30, 30)
        pdf.set_font("Helvetica", "", 10)
        _write_wrapped(pdf, analisis_texto)
        pdf.ln(2)

    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(14, 110, 79)
    pdf.cell(0, 8, "Recomendaciones", ln=1)
    pdf.set_text_color(30, 30, 30)
    pdf.set_font("Helvetica", "", 10)
    for r in recomendaciones:
        _write_wrapped(pdf, f"- {r}")
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(14, 110, 79)
    pdf.cell(0, 8, f"Detalle de acciones afirmativas (primeros 40 registros de {len(df)})", ln=1)
    pdf.set_text_color(30, 30, 30)
    pdf.set_font("Helvetica", "B", 8)
    encabezados = ["Fecha", "Nombre", "Localidad", "Tipo de accion", "Estado", "Presupuesto"]
    anchos = [20, 45, 30, 45, 25, 25]
    for h, w in zip(encabezados, anchos):
        pdf.cell(w, 6, h, border=1)
    pdf.ln()
    pdf.set_font("Helvetica", "", 7.5)
    for _, fila in df.head(40).iterrows():
        valores = [
            str(fila["fecha"].date()) if pd.notna(fila["fecha"]) else "",
            str(fila["nombre"])[:28],
            str(fila["localidad"])[:18],
            str(fila["tipo_accion"])[:28],
            str(fila["estado"])[:15],
            f"${fila['presupuesto']:,.0f}",
        ]
        for v, w in zip(valores, anchos):
            pdf.cell(w, 6, v, border=1)
        pdf.ln()

    return bytes(pdf.output())


# --------------------------------------------------------------------------------
# Word
# --------------------------------------------------------------------------------
def generar_reporte_word(df: pd.DataFrame, kpis: dict, recomendaciones: list, analisis_texto: str = "") -> bytes:
    doc = Document()

    titulo = doc.add_heading(f"{APP_NAME} — Reporte Ejecutivo", level=0)
    titulo.runs[0].font.color.rgb = RGBColor(0x0E, 0x6E, 0x4F)

    doc.add_paragraph(APP_FULL_NAME)
    doc.add_paragraph(APP_ENTITY)
    doc.add_paragraph(f"Generado el {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("Indicadores clave", level=1)
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Light Grid Accent 1"
    hdr = tabla.rows[0].cells
    hdr[0].text, hdr[1].text = "Indicador", "Valor"
    filas_kpi = [
        ("Recicladores registrados", kpis.get("total_recicladores", 0)),
        ("Organizaciones", kpis.get("total_organizaciones", 0)),
        ("Acciones afirmativas", kpis.get("total_acciones", 0)),
        ("Beneficiarios unicos", kpis.get("beneficiarios_unicos", 0)),
        ("Cobertura (%)", f"{kpis.get('cobertura_pct', 0)}%"),
        ("Presupuesto total", f"${kpis.get('presupuesto_total', 0):,.0f}"),
        ("Presupuesto ejecutado", f"${kpis.get('presupuesto_ejecutado', 0):,.0f}"),
        ("Acciones pendientes", kpis.get("acciones_pendientes", 0)),
        ("Alertas activas", kpis.get("alertas_activas", 0)),
    ]
    for etiqueta, valor in filas_kpi:
        fila = tabla.add_row().cells
        fila[0].text, fila[1].text = etiqueta, str(valor)

    if analisis_texto:
        doc.add_heading("Analisis ejecutivo", level=1)
        doc.add_paragraph(analisis_texto)

    doc.add_heading("Recomendaciones", level=1)
    for r in recomendaciones:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_heading(f"Detalle de acciones afirmativas (primeros 40 de {len(df)})", level=1)
    columnas = ["fecha", "nombre", "localidad", "tipo_accion", "estado", "presupuesto"]
    columnas = [c for c in columnas if c in df.columns]
    tabla2 = doc.add_table(rows=1, cols=len(columnas))
    tabla2.style = "Light List Accent 1"
    for i, c in enumerate(columnas):
        tabla2.rows[0].cells[i].text = c.replace("_", " ").title()
    for _, fila in df.head(40).iterrows():
        celdas = tabla2.add_row().cells
        for i, c in enumerate(columnas):
            valor = fila[c]
            if c == "fecha" and pd.notna(valor):
                valor = valor.strftime("%Y-%m-%d")
            celdas[i].text = str(valor)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# --------------------------------------------------------------------------------
# Informe Ejecutivo (PDF institucional para Alta Direccion)
# --------------------------------------------------------------------------------
# Sistema de diseno propio construido unicamente con primitivas nativas de
# fpdf2 (rect, ellipse, line, cell, image) para no depender de librerias de
# graficos adicionales (matplotlib/kaleido) ni de metodos de fpdf2 sin
# verificar (polygon, rotacion). El objetivo es una presentacion ejecutiva
# tipo firma consultora (Deloitte / McKinsey / Power BI): tarjetas KPI,
# medidor tipo gauge, dona, ranking con color de intensidad, tablero de
# alertas y matriz visual de recomendaciones. Toda la logica e informacion
# (KPIs, Valor Publico, alertas, recomendaciones) se mantiene sin cambios;
# solo cambia la forma en que se presenta.
# --------------------------------------------------------------------------------

_VERDE = (14, 110, 79)
_VERDE_OSCURO = (9, 74, 53)
_VERDE_EXITO = (30, 142, 90)
_VERDE_EXITO_CLARO = (224, 241, 232)
_AZUL = (36, 100, 180)
_GRIS_CLARO = (246, 247, 248)
_GRIS_MEDIO = (217, 222, 226)
_GRIS_OSCURO = (59, 65, 72)
_GRIS_TEXTO = (110, 116, 122)
_BLANCO = (255, 255, 255)
_AMARILLO = (224, 161, 0)
_AMARILLO_CLARO = (252, 244, 224)
_ROJO = (192, 57, 43)
_ROJO_CLARO = (250, 231, 228)

_SEMAFORO_COLORES = {"🟢": _VERDE_EXITO, "🟡": _AMARILLO, "🔴": _ROJO}
_PRIORIDAD_COLORES = {"Alta": _ROJO, "Media": _AMARILLO, "Baja": _VERDE_EXITO}
_SEVERIDAD_COLORES = {"alta": _ROJO, "media": _AMARILLO, "baja": _VERDE_EXITO}
_SEVERIDAD_COLORES_CLARO = {"alta": _ROJO_CLARO, "media": _AMARILLO_CLARO, "baja": _VERDE_EXITO_CLARO}


def _tinte(color: tuple, factor: float = 0.84) -> tuple:
    """Version clara (tinte) de un color RGB, usada como fondo de insignias
    y tarjetas para mantener contraste con el color solido del icono/texto."""
    return tuple(int(c + (255 - c) * factor) for c in color)


def _linea(pdf: FPDF, x1: float, y1: float, x2: float, y2: float, color: tuple, grosor: float = 0.3):
    pdf.set_draw_color(*color)
    pdf.set_line_width(grosor)
    pdf.line(x1, y1, x2, y2)


def _circulo(pdf: FPDF, cx: float, cy: float, r: float, color: tuple):
    pdf.set_fill_color(*color)
    pdf.ellipse(cx - r, cy - r, r * 2, r * 2, style="F")


def _arco_grueso(pdf: FPDF, cx: float, cy: float, r: float, angulo_ini: float, angulo_fin: float,
                  color: tuple, grosor: float):
    """
    Dibuja una banda de arco grueso (usada en el gauge y en la dona) uniendo
    muchos segmentos cortos de linea con un ancho de trazo amplio. Se evita
    deliberadamente cualquier metodo de arco/rotacion de fpdf2 cuya
    disponibilidad exacta en la version instalada no se pudo verificar.
    """
    pdf.set_draw_color(*color)
    pdf.set_line_width(grosor)
    span = angulo_fin - angulo_ini
    n = max(int(abs(span) / 4), 6)
    for i in range(n):
        a0 = math.radians(angulo_ini + span * i / n)
        a1 = math.radians(angulo_ini + span * (i + 1) / n)
        x0, y0 = cx + r * math.cos(a0), cy + r * math.sin(a0)
        x1, y1 = cx + r * math.cos(a1), cy + r * math.sin(a1)
        pdf.line(x0, y0, x1, y1)


def _rect_suave(pdf: FPDF, x: float, y: float, w: float, h: float, color: tuple = None,
                 borde: tuple = None, grosor_borde: float = 0.25, radio: float = 2.2):
    """
    Dibuja un rectangulo (tarjeta) con esquinas redondeadas si la version de
    fpdf2 instalada lo soporta (round_corners/corner_radius); si no, recurre
    de forma silenciosa a un rectangulo recto, para que el informe nunca
    falle por una diferencia de version.
    """
    estilo = ""
    if color:
        pdf.set_fill_color(*color)
        estilo += "F"
    if borde:
        pdf.set_draw_color(*borde)
        pdf.set_line_width(grosor_borde)
        estilo += "D"
    if not estilo:
        return
    try:
        pdf.rect(x, y, w, h, style=estilo, round_corners=True, corner_radius=radio)
    except Exception:
        pdf.rect(x, y, w, h, style=estilo)


def _asegurar_espacio(pdf: FPDF, alto_necesario: float):
    """
    Fuerza un salto de pagina manual si el bloque que sigue (dibujado con
    rect()/ellipse()/line(), que no activan el salto de pagina automatico de
    fpdf2) no cabe en el espacio restante de la pagina actual.
    """
    espacio_restante = pdf.h - pdf.b_margin - pdf.get_y()
    if alto_necesario > espacio_restante:
        pdf.add_page()


def _texto_multilinea(pdf: FPDF, x: float, y: float, w: float, texto: str, tam: float = 8.5,
                       alto_linea: float = 4.2, color: tuple = _GRIS_OSCURO, negrita: bool = False,
                       max_lineas: int = None) -> float:
    """Envuelve y dibuja texto dentro de un ancho `w` (en mm), linea por
    linea, usando cell() (ver `_write_wrapped` para el motivo de evitar
    multi_cell()). Retorna la coordenada Y siguiente al bloque de texto."""
    pdf.set_font("Helvetica", "B" if negrita else "", tam)
    pdf.set_text_color(*color)
    ancho_caracteres = max(int(w / (tam * 0.235)), 10)
    lineas = textwrap.wrap(_pdf_safe_text(texto), width=ancho_caracteres) or [""]
    if max_lineas:
        lineas = lineas[:max_lineas]
    for i, linea in enumerate(lineas):
        pdf.set_xy(x, y + i * alto_linea)
        pdf.cell(w, alto_linea, linea, ln=0)
    return y + len(lineas) * alto_linea


# --------------------------------------------------------------------------------
# Iconografia lineal minimalista (dibujada solo con line()/ellipse())
# --------------------------------------------------------------------------------
def _icono(pdf: FPDF, nombre: str, x: float, y: float, s: float, color: tuple):
    g = max(s * 0.11, 0.45)
    cx, cy = x + s / 2, y + s / 2
    if nombre == "reciclaje":
        for a0 in (18, 138, 258):
            _arco_grueso(pdf, cx, cy, s * 0.36, a0, a0 + 76, color, g * 1.5)
    elif nombre == "personas":
        _circulo(pdf, cx, y + s * 0.26, s * 0.15, color)
        pdf.set_fill_color(*color)
        pdf.ellipse(x + s * 0.16, y + s * 0.48, s * 0.68, s * 0.42, style="F")
    elif nombre == "documento":
        _linea(pdf, x + s * 0.24, y + s * 0.08, x + s * 0.24, y + s * 0.92, color, g)
        _linea(pdf, x + s * 0.24, y + s * 0.08, x + s * 0.76, y + s * 0.08, color, g)
        _linea(pdf, x + s * 0.76, y + s * 0.08, x + s * 0.76, y + s * 0.92, color, g)
        _linea(pdf, x + s * 0.24, y + s * 0.92, x + s * 0.76, y + s * 0.92, color, g)
        for i in range(3):
            yy = y + s * 0.3 + i * s * 0.2
            _linea(pdf, x + s * 0.34, yy, x + s * 0.66, yy, color, g * 0.7)
    elif nombre == "dinero":
        _circulo(pdf, cx, cy, s * 0.38, color)
        pdf.set_fill_color(255, 255, 255)
        pdf.ellipse(cx - s * 0.24, cy - s * 0.24, s * 0.48, s * 0.48, style="F")
        _linea(pdf, cx, cy - s * 0.16, cx, cy + s * 0.16, color, g)
    elif nombre == "objetivo":
        _circulo(pdf, cx, cy, s * 0.4, color)
        pdf.set_fill_color(255, 255, 255)
        pdf.ellipse(cx - s * 0.26, cy - s * 0.26, s * 0.52, s * 0.52, style="F")
        _circulo(pdf, cx, cy, s * 0.13, color)
    elif nombre == "reloj":
        _circulo(pdf, cx, cy, s * 0.4, color)
        pdf.set_fill_color(255, 255, 255)
        pdf.ellipse(cx - s * 0.32, cy - s * 0.32, s * 0.64, s * 0.64, style="F")
        _linea(pdf, cx, cy, cx, cy - s * 0.22, color, g)
        _linea(pdf, cx, cy, cx + s * 0.16, cy, color, g)
    elif nombre == "check":
        _linea(pdf, x + s * 0.2, y + s * 0.52, x + s * 0.42, y + s * 0.74, color, g * 1.4)
        _linea(pdf, x + s * 0.42, y + s * 0.74, x + s * 0.82, y + s * 0.26, color, g * 1.4)
    elif nombre == "pin":
        _circulo(pdf, cx, y + s * 0.34, s * 0.24, color)
        pdf.set_fill_color(255, 255, 255)
        pdf.ellipse(cx - s * 0.09, y + s * 0.25, s * 0.18, s * 0.18, style="F")
        _linea(pdf, cx, y + s * 0.34, cx, y + s * 0.9, color, g * 1.3)
    elif nombre == "bombilla":
        _circulo(pdf, cx, y + s * 0.36, s * 0.28, color)
        pdf.set_fill_color(255, 255, 255)
        pdf.ellipse(cx - s * 0.1, y + s * 0.6, s * 0.2, s * 0.1, style="F")
        _linea(pdf, cx - s * 0.1, y + s * 0.74, cx + s * 0.1, y + s * 0.74, color, g)
        _linea(pdf, cx - s * 0.08, y + s * 0.84, cx + s * 0.08, y + s * 0.84, color, g)
    elif nombre == "alerta":
        pdf.set_draw_color(*color)
        pdf.set_line_width(g * 1.3)
        pdf.line(cx, y + s * 0.06, x + s * 0.92, y + s * 0.88)
        pdf.line(x + s * 0.92, y + s * 0.88, x + s * 0.08, y + s * 0.88)
        pdf.line(x + s * 0.08, y + s * 0.88, cx, y + s * 0.06)
        _linea(pdf, cx, y + s * 0.36, cx, y + s * 0.62, color, g)
        _circulo(pdf, cx, y + s * 0.76, g * 0.55, color)
    elif nombre == "tendencia":
        pdf.set_fill_color(*color)
        alturas = (0.3, 0.5, 0.4, 0.7, 0.9)
        anchobar = s * 0.14
        for i, hfrac in enumerate(alturas):
            bx = x + s * 0.05 + i * (anchobar + s * 0.04)
            bh = s * hfrac
            pdf.rect(bx, y + s - bh, anchobar, bh, "F")
    elif nombre == "mapa":
        _linea(pdf, x + s * 0.1, y + s * 0.3, x + s * 0.9, y + s * 0.3, color, g * 0.6)
        _linea(pdf, x + s * 0.1, y + s * 0.6, x + s * 0.9, y + s * 0.6, color, g * 0.6)
        _linea(pdf, x + s * 0.35, y + s * 0.1, x + s * 0.35, y + s * 0.9, color, g * 0.6)
        _linea(pdf, x + s * 0.65, y + s * 0.1, x + s * 0.65, y + s * 0.9, color, g * 0.6)
        _circulo(pdf, x + s * 0.65, y + s * 0.3, s * 0.09, color)
    elif nombre == "escudo":
        pdf.set_draw_color(*color)
        pdf.set_line_width(g)
        pdf.line(x + s * 0.2, y + s * 0.12, x + s * 0.8, y + s * 0.12)
        pdf.line(x + s * 0.2, y + s * 0.12, x + s * 0.2, y + s * 0.55)
        pdf.line(x + s * 0.8, y + s * 0.12, x + s * 0.8, y + s * 0.55)
        pdf.line(x + s * 0.2, y + s * 0.55, cx, y + s * 0.9)
        pdf.line(x + s * 0.8, y + s * 0.55, cx, y + s * 0.9)


def _insignia_icono(pdf: FPDF, nombre: str, x: float, y: float, s: float, color_fondo: tuple, color_icono: tuple):
    """Insignia circular/redondeada de fondo con un icono lineal dentro,
    usada en tarjetas KPI, bloques narrativos, alertas y recomendaciones."""
    _rect_suave(pdf, x, y, s, s, color=color_fondo, radio=s * 0.24)
    margen = s * 0.2
    _icono(pdf, nombre, x + margen, y + margen, s - margen * 2, color_icono)


def _ilustracion_economia_circular(pdf: FPDF, cx: float, cy: float, r: float):
    """Ilustracion de portada/cierre: economia circular, reciclaje e
    innovacion (aro exterior blanco + tres arcos + nucleo verde)."""
    _circulo(pdf, cx, cy, r, _BLANCO)
    for a0 in (10, 130, 250):
        _arco_grueso(pdf, cx, cy, r * 0.62, a0, a0 + 82, _VERDE, r * 0.11)
    _circulo(pdf, cx, cy, r * 0.22, _VERDE)


def _color_intensidad(valor: float, maximo: float, color_base: tuple = _VERDE) -> tuple:
    """Interpola entre una version clara y el color base segun la
    proporcion del valor sobre el maximo (mapa de calor por intensidad)."""
    ratio = (valor / maximo) if maximo else 0
    ratio = max(0.18, min(ratio, 1))
    claro = _tinte(color_base, 0.82)
    return tuple(int(claro[i] + (color_base[i] - claro[i]) * ratio) for i in range(3))


# --------------------------------------------------------------------------------
# Componentes visuales compuestos
# --------------------------------------------------------------------------------
def _gauge_valor_publico(pdf: FPDF, x: float, y: float, radio: float, pct: float, color: tuple):
    """Medidor tipo gauge (semicirculo) para el Valor Publico Generado."""
    pct = max(0, min(pct, 100))
    _arco_grueso(pdf, x, y, radio, 180, 360, _GRIS_MEDIO, radio * 0.22)
    _arco_grueso(pdf, x, y, radio, 180, 180 + 180 * pct / 100, color, radio * 0.22)
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(*_GRIS_OSCURO)
    pdf.set_xy(x - radio, y - radio * 0.32)
    pdf.cell(radio * 2, 10, f"{pct:.0f}%", align="C")


def _dona(pdf: FPDF, cx: float, cy: float, r_ext: float, grosor: float, segmentos: list):
    """segmentos: lista de tuplas (etiqueta, valor, color). Dibuja una dona
    (grafico circular con centro hueco) proporcional a los valores."""
    total = sum(v for _, v, _ in segmentos) or 1
    angulo = -90
    for _, valor, color in segmentos:
        span = valor / total * 360
        if span <= 0:
            continue
        _arco_grueso(pdf, cx, cy, r_ext, angulo, angulo + span, color, grosor)
        angulo += span


def _tarjeta_kpi(pdf: FPDF, x: float, y: float, w: float, h: float, icono: str, color: tuple,
                  etiqueta: str, valor, nota: str = ""):
    _rect_suave(pdf, x, y, w, h, color=_BLANCO, borde=_GRIS_MEDIO, grosor_borde=0.25, radio=2.5)
    _insignia_icono(pdf, icono, x + 4, y + 4, 8, _tinte(color), color)
    pdf.set_xy(x + 4, y + 13.5)
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(*_GRIS_OSCURO)
    pdf.cell(w - 8, 7, _pdf_safe_text(str(valor)), ln=0)
    _texto_multilinea(pdf, x + 4, y + 21.5, w - 8, etiqueta, tam=7.3, alto_linea=3.3,
                       color=_GRIS_TEXTO, max_lineas=2)
    if nota:
        color_nota = _VERDE_EXITO if nota.startswith("+") else (_ROJO if nota.startswith("-") else _GRIS_TEXTO)
        pdf.set_xy(x + 4, y + h - 5.5)
        pdf.set_font("Helvetica", "B", 6.8)
        pdf.set_text_color(*color_nota)
        pdf.cell(w - 8, 4, _pdf_safe_text(nota)[:34], ln=0)


def _fila_progreso(pdf: FPDF, x: float, y: float, w: float, etiqueta: str, pct: float, color: tuple,
                    alto: float = 5.4):
    pct = max(0, min(pct, 100))
    pdf.set_xy(x, y)
    pdf.set_font("Helvetica", "", 8.3)
    pdf.set_text_color(*_GRIS_OSCURO)
    ancho_etq = w * 0.4
    pdf.cell(ancho_etq, alto, _pdf_safe_text(etiqueta), ln=0)
    barra_x = x + ancho_etq + 2
    barra_w = w * 0.44
    _rect_suave(pdf, barra_x, y + 0.9, barra_w, alto - 1.8, color=_GRIS_CLARO, radio=1)
    _rect_suave(pdf, barra_x, y + 0.9, barra_w * pct / 100, alto - 1.8, color=color, radio=1)
    pdf.set_xy(barra_x + barra_w + 2, y)
    pdf.set_font("Helvetica", "B", 8.2)
    pdf.set_text_color(*_GRIS_OSCURO)
    pdf.cell(w - ancho_etq - barra_w - 4, alto, f"{pct:.0f}%", ln=0)


def _fila_barra_moderna(pdf: FPDF, x: float, y: float, w: float, etiqueta: str, valor: float,
                         valor_max: float, color: tuple, formato: str = "{:,.0f}", alto: float = 6.4):
    pdf.set_xy(x, y)
    pdf.set_font("Helvetica", "", 8.2)
    pdf.set_text_color(*_GRIS_OSCURO)
    ancho_etq = w * 0.36
    pdf.cell(ancho_etq, alto, _pdf_safe_text(str(etiqueta))[:32], ln=0)
    barra_x = x + ancho_etq + 2
    barra_w_total = w * 0.42
    barra_w = (valor / valor_max * barra_w_total) if valor_max else 0
    _rect_suave(pdf, barra_x, y + 1.1, barra_w_total, alto - 2.2, color=_GRIS_CLARO, radio=1)
    _rect_suave(pdf, barra_x, y + 1.1, max(barra_w, 1.2), alto - 2.2, color=color, radio=1)
    pdf.set_xy(barra_x + barra_w_total + 2, y)
    pdf.set_font("Helvetica", "B", 8)
    pdf.set_text_color(*_GRIS_OSCURO)
    pdf.cell(w - ancho_etq - barra_w_total - 4, alto, formato.format(valor), ln=0)


def _tarjeta_alerta(pdf: FPDF, x: float, y: float, w: float, alerta: dict) -> float:
    severidad = alerta.get("severidad", "media")
    color = _SEVERIDAD_COLORES.get(severidad, _AMARILLO)
    color_claro = _SEVERIDAD_COLORES_CLARO.get(severidad, _AMARILLO_CLARO)
    h = 21
    _rect_suave(pdf, x, y, w, h, color=color_claro, radio=2.2)
    _rect_suave(pdf, x, y, 2.2, h, color=color, radio=0)
    _insignia_icono(pdf, "alerta", x + 5, y + 4, 8, _BLANCO, color)
    pdf.set_xy(x + 16, y + 3.5)
    pdf.set_font("Helvetica", "B", 8.5)
    pdf.set_text_color(*color)
    pdf.cell(w - 20, 5, _pdf_safe_text(f"{severidad.upper()} - {alerta.get('tipo', '')}")[:60], ln=0)
    _texto_multilinea(pdf, x + 16, y + 9.5, w - 22, alerta.get("mensaje", ""), tam=7.5, alto_linea=3.3,
                       color=_GRIS_OSCURO, max_lineas=2)
    return h


def _fila_matriz_recomendacion(pdf: FPDF, x: float, y: float, w: float, item: dict) -> float:
    h = 19
    _rect_suave(pdf, x, y, w, h, color=_BLANCO, borde=_GRIS_MEDIO, grosor_borde=0.2, radio=1.6)
    color_pr = _PRIORIDAD_COLORES.get(item.get("prioridad", "Media"), _AMARILLO)
    _insignia_icono(pdf, item.get("icono", "objetivo"), x + 3, y + 3, 8, _tinte(color_pr), color_pr)

    col_texto_x = x + 14
    col_texto_w = w * 0.42
    _texto_multilinea(pdf, col_texto_x, y + 2.3, col_texto_w, item.get("texto", ""), tam=7.4,
                       alto_linea=3.1, color=_GRIS_OSCURO, max_lineas=4)

    col2_x = col_texto_x + col_texto_w + 2
    col2_w = w * 0.2
    _texto_multilinea(pdf, col2_x, y + 2.3, col2_w, item.get("impacto", ""), tam=7, alto_linea=3.1,
                       color=_GRIS_TEXTO, max_lineas=3)

    col3_x = col2_x + col2_w + 2
    _rect_suave(pdf, col3_x, y + 2.8, 16, 5.5, color=color_pr, radio=1.4)
    pdf.set_xy(col3_x, y + 3.7)
    pdf.set_font("Helvetica", "B", 6.6)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(16, 4, _pdf_safe_text(item.get("prioridad", "")), align="C")

    col4_x = col3_x + 19
    col4_w = max(w - (col4_x - x) - 2, 10)
    _texto_multilinea(pdf, col4_x, y + 2.3, col4_w,
                       f"{item.get('responsable', '')} | {item.get('horizonte', '')}",
                       tam=6.6, alto_linea=3, color=_GRIS_TEXTO, max_lineas=3)
    return h


def _titulo_pagina(pdf: FPDF, titulo: str, subtitulo: str = ""):
    pdf.set_xy(10, pdf.get_y())
    pdf.set_font("Helvetica", "B", 15)
    pdf.set_text_color(*_VERDE)
    pdf.cell(0, 9, _pdf_safe_text(titulo), ln=1)
    if subtitulo:
        pdf.set_x(10)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*_GRIS_TEXTO)
        pdf.cell(0, 5, _pdf_safe_text(subtitulo), ln=1)
    _linea(pdf, 10, pdf.get_y() + 1, 200, pdf.get_y() + 1, _GRIS_MEDIO, 0.4)
    pdf.ln(5)


class InformeEjecutivoPDF(FPDF):
    """
    Documento PDF de informe ejecutivo para Alta Direccion. A diferencia de
    `ReportePDF` (reporte operativo de la pagina Reportes), este documento
    incluye una portada institucional propia; por eso el banner verde de
    encabezado solo se dibuja a partir de la segunda pagina.
    """

    def header(self):
        if self.page_no() == 1:
            return  # La portada dibuja su propio encabezado institucional.
        self.set_fill_color(*_VERDE)
        self.rect(0, 0, 210, 16, "F")
        self.set_fill_color(*_VERDE_OSCURO)
        self.rect(0, 16, 210, 1.2, "F")
        self.set_text_color(255, 255, 255)
        self.set_font("Helvetica", "B", 11)
        self.set_xy(10, 4)
        self.cell(130, 8, _pdf_safe_text(f"{APP_NAME} - Informe Ejecutivo"), ln=0)
        self.set_font("Helvetica", "", 8)
        self.set_xy(130, 5)
        self.cell(70, 8, _pdf_safe_text(fecha_larga_es()), align="R")
        self.set_xy(10, 21)
        self.set_text_color(30, 30, 30)

    def footer(self):
        self.set_y(-14)
        self.set_draw_color(*_GRIS_MEDIO)
        self.set_line_width(0.2)
        self.line(10, self.get_y(), 200, self.get_y())
        self.set_font("Helvetica", "I", 7.5)
        self.set_text_color(140, 140, 140)
        self.set_xy(10, self.get_y() + 1.5)
        self.cell(120, 6, _pdf_safe_text(f"{APP_NAME} v{APP_VERSION} | {APP_ENTITY}"), ln=0)
        self.set_xy(140, self.get_y())
        self.cell(60, 6, f"Pagina {self.page_no()}", align="R")


# --------------------------------------------------------------------------------
# Paginas del Informe Ejecutivo
# --------------------------------------------------------------------------------
def _pagina_portada(pdf: FPDF, usuario: dict):
    pdf.add_page()
    pdf.set_fill_color(*_VERDE_OSCURO)
    pdf.rect(0, 0, 210, 297, "F")
    pdf.set_fill_color(*_VERDE)
    pdf.rect(0, 0, 210, 100, "F")
    pdf.set_fill_color(*_AZUL)
    pdf.rect(0, 100, 210, 2, "F")

    try:
        if LOGO_PATH.exists():
            pdf.image(str(LOGO_PATH), x=15, y=14, w=30)
    except Exception:
        pass  # La portada se muestra igual si el logo no esta disponible.

    pdf.set_left_margin(15)
    pdf.set_right_margin(15)
    pdf.set_text_color(255, 255, 255)
    pdf.set_xy(15, 50)
    pdf.set_font("Helvetica", "B", 30)
    pdf.cell(0, 12, "SIGATA", ln=1)
    pdf.set_x(15)
    pdf.set_font("Helvetica", "", 11.5)
    _write_wrapped(pdf, APP_FULL_NAME, ancho_caracteres=54, alto_linea=5.6)
    pdf.set_x(15)
    pdf.ln(1)
    pdf.set_x(15)
    pdf.set_font("Helvetica", "B", 15)
    pdf.cell(0, 9, "Informe Ejecutivo", ln=1)

    _ilustracion_economia_circular(pdf, cx=175, cy=44, r=25)

    y_badges = 112
    badges = [
        ("Proyecto de Inversion", f"{PROYECTO_INVERSION['codigo']} - {PROYECTO_INVERSION['nombre']}"),
        ("Objetivo ODS", "ODS 12: Produccion y Consumo Responsables"),
        ("Dependencia responsable", APP_ENTITY),
    ]
    yb = y_badges
    for titulo_b, texto_b in badges:
        _rect_suave(pdf, 15, yb, 180, 16, color=_BLANCO, radio=2.4)
        pdf.set_xy(20, yb + 2.2)
        pdf.set_font("Helvetica", "B", 7.6)
        pdf.set_text_color(*_VERDE)
        pdf.cell(0, 4.2, _pdf_safe_text(titulo_b.upper()), ln=1)
        pdf.set_x(20)
        pdf.set_font("Helvetica", "", 9.5)
        pdf.set_text_color(*_GRIS_OSCURO)
        pdf.cell(0, 5.2, _pdf_safe_text(texto_b)[:78], ln=1)
        yb += 20

    y_ico = yb + 8
    pdf.set_xy(15, y_ico - 6)
    pdf.set_font("Helvetica", "B", 8.5)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 5, "ECONOMIA CIRCULAR  |  GESTION PUBLICA  |  DATOS  |  INNOVACION", ln=1)
    iconos_pie = ["reciclaje", "escudo", "tendencia", "bombilla"]
    ancho_ico = 180 / 4
    for i, ic in enumerate(iconos_pie):
        cx_i = 15 + ancho_ico * i + ancho_ico / 2
        _circulo(pdf, cx_i, y_ico + 10, 8, _BLANCO)
        _icono(pdf, ic, cx_i - 4.4, y_ico + 5.6, 8.8, _VERDE)

    pdf.set_draw_color(255, 255, 255)
    pdf.set_line_width(0.3)
    pdf.line(15, 246, 195, 246)
    pdf.set_font("Helvetica", "B", 8.5)
    pdf.set_text_color(255, 255, 255)
    pdf.set_xy(15, 250)
    pdf.cell(60, 5, "FECHA DEL INFORME", ln=0)
    pdf.cell(60, 5, "GENERADO POR", ln=0)
    pdf.cell(60, 5, "CARGO", ln=1)
    pdf.set_font("Helvetica", "", 9.3)
    pdf.set_x(15)
    pdf.cell(60, 6, _pdf_safe_text(fecha_larga_es())[:32], ln=0)
    pdf.cell(60, 6, _pdf_safe_text(usuario["nombre"])[:32], ln=0)
    pdf.cell(60, 6, _pdf_safe_text(usuario["cargo"])[:32], ln=1)

    pdf.set_xy(15, 268)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(205, 224, 216)
    _write_wrapped(
        pdf,
        "Documento de circulacion institucional preparado para el Comite Institucional de Gestion y "
        "Desempeno, Alta Direccion de la UAESP y entidades de control.",
        ancho_caracteres=92, alto_linea=4.2,
    )


def _pagina_dashboard_kpi(pdf: FPDF, kpi: dict, variacion: dict):
    pdf.add_page()
    _titulo_pagina(pdf, "Indicadores Estrategicos",
                    "Panorama general de la gestion de acciones afirmativas en el periodo y filtros seleccionados")

    nota_acciones = ""
    if variacion.get("direccion") == "up":
        nota_acciones = f"+{variacion['delta_pct']:.1f}% vs. periodo anterior"
    elif variacion.get("direccion") == "down":
        nota_acciones = f"-{abs(variacion['delta_pct']):.1f}% vs. periodo anterior"
    elif variacion.get("direccion") == "flat":
        nota_acciones = "estable vs. periodo anterior"

    color_alertas = _ROJO if kpi.get("alertas_activas", 0) > 3 else _AMARILLO

    tarjetas = [
        ("personas", _VERDE, "Recicladores registrados", f"{kpi.get('total_recicladores', 0):,}", ""),
        ("objetivo", _AZUL, "Organizaciones (ORO)", f"{kpi.get('total_organizaciones', 0):,}", ""),
        ("documento", _VERDE, "Acciones afirmativas", f"{kpi.get('total_acciones', 0):,}", nota_acciones),
        ("personas", _AZUL, "Beneficiarios unicos", f"{kpi.get('beneficiarios_unicos', 0):,}", ""),
        ("tendencia", _VERDE, "Cobertura", f"{kpi.get('cobertura_pct', 0):.1f}%", "sobre poblacion registrada"),
        ("dinero", _VERDE, "Presupuesto ejecutado", f"${kpi.get('presupuesto_ejecutado', 0):,.0f}",
         f"{kpi.get('pct_ejecucion_presupuestal', 0):.1f}% del total"),
        ("dinero", _AZUL, "Presupuesto disponible", f"${kpi.get('presupuesto_disponible', 0):,.0f}", ""),
        ("check", _VERDE, "Ejecutadas este ano", f"{kpi.get('acciones_ejecutadas_anio', 0):,}", ""),
        ("reloj", _AMARILLO, "Acciones pendientes", f"{kpi.get('acciones_pendientes', 0):,}", ""),
        ("alerta", color_alertas, "Alertas activas", f"{kpi.get('alertas_activas', 0):,}",
         "revisar pagina de Alertas"),
    ]

    col_w = (190 - 10) / 3
    fila_h = 34
    filas_totales = math.ceil(len(tarjetas) / 3)
    _asegurar_espacio(pdf, filas_totales * (fila_h + 5))
    x0, y0 = 10, pdf.get_y()
    for i, (icono, color, etiqueta, valor, nota) in enumerate(tarjetas):
        fila, col = divmod(i, 3)
        x = x0 + col * (col_w + 5)
        y = y0 + fila * (fila_h + 5)
        _tarjeta_kpi(pdf, x, y, col_w, fila_h, icono, color, etiqueta, valor, nota)
    pdf.set_y(y0 + filas_totales * (fila_h + 5) + 3)


def _pagina_valor_publico(pdf: FPDF, valor_publico: dict):
    pdf.add_page()
    _titulo_pagina(pdf, "Valor Publico Generado",
                    "Indicador sintetico de impacto institucional (no corresponde a una cifra financiera)")

    color_sem = _SEMAFORO_COLORES.get(valor_publico.get("semaforo"), _AMARILLO)
    y0 = pdf.get_y()
    _rect_suave(pdf, 10, y0, 190, 62, color=_GRIS_CLARO, radio=3)
    _gauge_valor_publico(pdf, x=55, y=y0 + 42, radio=30, pct=valor_publico.get("valor_pct", 0), color=color_sem)

    pdf.set_xy(100, y0 + 10)
    pdf.set_font("Helvetica", "B", 15)
    pdf.set_text_color(*color_sem)
    pdf.cell(90, 8, _pdf_safe_text(valor_publico.get("interpretacion", "")), ln=1)

    _circulo(pdf, 104, y0 + 23.5, 2.6, color_sem)
    pdf.set_xy(110, y0 + 20.5)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(*_GRIS_OSCURO)
    pdf.cell(80, 6, _pdf_safe_text(valor_publico.get("tendencia_texto", ""))[:55], ln=1)

    _texto_multilinea(
        pdf, 100, y0 + 30, 90,
        "Combina cobertura de beneficiarios, trazabilidad, completitud de registros, evidencia "
        "documental y oportunidad del seguimiento.",
        tam=8, alto_linea=3.6, color=_GRIS_TEXTO, max_lineas=4,
    )

    pdf.set_y(y0 + 68)
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(*_VERDE)
    pdf.cell(0, 7, "Componentes del indicador", ln=1)
    pdf.ln(1)

    componentes = valor_publico.get("componentes", {})
    etiquetas_comp = {
        "cobertura": "Cobertura de beneficiarios",
        "trazabilidad": "Trazabilidad de acciones",
        "completitud_registros": "Completitud de registros",
        "evidencia_documental": "Evidencia documental",
        "seguimiento_oportuno": "Seguimiento oportuno",
    }
    y = pdf.get_y()
    for clave, etiqueta in etiquetas_comp.items():
        valor_c = componentes.get(clave, 0)
        color_c = _VERDE_EXITO if valor_c >= 65 else (_AMARILLO if valor_c >= 40 else _ROJO)
        _fila_progreso(pdf, 10, y, 190, etiqueta, valor_c, color_c)
        y += 7.6
    pdf.set_y(y + 4)


def _pagina_resumen_ejecutivo(pdf: FPDF, hallazgos: dict):
    pdf.add_page()
    _titulo_pagina(pdf, "Resumen Ejecutivo",
                    "Sintesis narrativa de la gestion en el periodo y filtros seleccionados")

    bloques = [
        ("tendencia", _AZUL, "Hallazgos principales", hallazgos.get("hallazgos", [])),
        ("check", _VERDE_EXITO, "Logros", hallazgos.get("logros", [])),
        ("bombilla", _AMARILLO, "Oportunidades", hallazgos.get("oportunidades", [])),
        ("alerta", _ROJO, "Riesgos", hallazgos.get("riesgos", [])),
        ("objetivo", _VERDE, "Acciones prioritarias", hallazgos.get("acciones_prioritarias", [])),
    ]
    for icono, color, titulo_b, items in bloques:
        alto = 12 + max(len(items), 1) * 7
        _asegurar_espacio(pdf, alto + 4)
        y = pdf.get_y()
        _rect_suave(pdf, 10, y, 190, alto, color=_tinte(color), radio=2.5)
        _insignia_icono(pdf, icono, 14, y + 4, 9, _BLANCO, color)
        pdf.set_xy(27, y + 4)
        pdf.set_font("Helvetica", "B", 10.5)
        pdf.set_text_color(*color)
        pdf.cell(0, 6, _pdf_safe_text(titulo_b), ln=1)
        yy = y + 11
        for item in items:
            _circulo(pdf, 30, yy + 1.6, 0.9, color)
            _texto_multilinea(pdf, 34, yy, 160, item, tam=8.3, alto_linea=3.6, color=_GRIS_OSCURO, max_lineas=2)
            yy += 7
        pdf.set_y(y + alto + 4)


def _pagina_graficos(pdf: FPDF, df: pd.DataFrame):
    pdf.add_page()
    _titulo_pagina(pdf, "Graficos Principales",
                    "Distribucion de acciones afirmativas por tipo, beneficio y estacionalidad")

    y0 = pdf.get_y()
    pdf.set_xy(10, y0)
    pdf.set_font("Helvetica", "B", 10.5)
    pdf.set_text_color(*_VERDE)
    pdf.cell(90, 6, "Acciones afirmativas por tipo", ln=0)
    pdf.set_xy(110, y0)
    pdf.cell(90, 6, "Distribucion de beneficios (Top 5)", ln=1)
    y = y0 + 8

    if "tipo_accion" in df.columns and not df.empty:
        conteo_tipo = df["tipo_accion"].value_counts().head(6)
        if not conteo_tipo.empty:
            max_val = conteo_tipo.max()
            yy = y
            for etiqueta, valor in conteo_tipo.items():
                _fila_barra_moderna(pdf, 10, yy, 90, str(etiqueta), valor, max_val, _VERDE)
                yy += 8

    if "beneficio" in df.columns and not df.empty:
        conteo_ben = df["beneficio"].value_counts()
        conteo_ben = conteo_ben[conteo_ben.index != ""].head(5)
        if not conteo_ben.empty:
            colores_dona = [_VERDE, _AZUL, _AMARILLO, (94, 174, 140), _GRIS_OSCURO]
            segmentos = [(str(et), val, colores_dona[i % len(colores_dona)])
                         for i, (et, val) in enumerate(conteo_ben.items())]
            _dona(pdf, cx=150, cy=y + 24, r_ext=18, grosor=8, segmentos=segmentos)
            total_ben = sum(v for _, v, _ in segmentos)
            pdf.set_xy(130, y + 21)
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(*_GRIS_OSCURO)
            pdf.cell(40, 6, f"{int(total_ben):,}", align="C")
            for i, (etq, val, color) in enumerate(segmentos):
                ly = y + 48 + i * 5.4
                _circulo(pdf, 112, ly + 1.4, 1.3, color)
                pdf.set_xy(116, ly)
                pdf.set_font("Helvetica", "", 7.2)
                pdf.set_text_color(*_GRIS_OSCURO)
                pdf.cell(78, 4, _pdf_safe_text(f"{etq[:26]} ({val})"), ln=0)

    y_evol = y0 + 100
    _asegurar_espacio(pdf, 55)
    pdf.set_xy(10, y_evol)
    pdf.set_font("Helvetica", "B", 10.5)
    pdf.set_text_color(*_VERDE)
    pdf.cell(0, 6, "Evolucion mensual (estacionalidad)", ln=1)

    if "fecha" in df.columns and not df.empty:
        d = df.dropna(subset=["fecha"]).copy()
        if not d.empty:
            meses = ["Ene", "Feb", "Mar", "Abr", "May", "Jun", "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]
            d["mes"] = d["fecha"].dt.month
            conteo_mes = d.groupby("mes").size().reindex(range(1, 13), fill_value=0)
            max_mes = max(conteo_mes.max(), 1)
            base_y = y_evol + 38
            ancho_barra = 12
            x_ini = 12
            for i, m in enumerate(range(1, 13)):
                v = conteo_mes.loc[m]
                alto_barra = (v / max_mes) * 28
                bx = x_ini + i * 15
                _rect_suave(pdf, bx, base_y - alto_barra, ancho_barra, max(alto_barra, 0.6), color=_VERDE, radio=1)
                pdf.set_xy(bx - 1, base_y + 1)
                pdf.set_font("Helvetica", "", 6.5)
                pdf.set_text_color(*_GRIS_TEXTO)
                pdf.cell(ancho_barra + 2, 4, meses[i], align="C")
    pdf.set_y(y_evol + 48)


def _pagina_cobertura(pdf: FPDF, df: pd.DataFrame):
    pdf.add_page()
    _titulo_pagina(pdf, "Cobertura Territorial",
                    "Beneficiarios unicos por localidad. El mapa interactivo completo esta disponible en la "
                    "aplicacion, pagina Mapa.")

    if "localidad" not in df.columns or df.empty:
        _texto_multilinea(pdf, 10, pdf.get_y(), 190,
                           "No hay informacion de cobertura territorial disponible para el periodo y filtros "
                           "seleccionados.", tam=9)
        return

    conteo = df.groupby("localidad")["documento"].nunique().sort_values(ascending=False)
    conteo = conteo[conteo.index != ""]
    if conteo.empty:
        _texto_multilinea(pdf, 10, pdf.get_y(), 190,
                           "No hay localidades con informacion de cobertura registrada.", tam=9)
        return

    total = conteo.sum()
    max_val = conteo.max()

    pdf.set_font("Helvetica", "B", 10.5)
    pdf.set_text_color(*_VERDE)
    pdf.cell(0, 6, f"Ranking de localidades ({len(conteo)} con cobertura registrada)", ln=1)
    y = pdf.get_y() + 2
    for loc, val in conteo.items():
        _asegurar_espacio(pdf, 8)
        y = pdf.get_y()
        pct = val / total * 100 if total else 0
        color = _color_intensidad(val, max_val, _VERDE)
        _fila_barra_moderna(pdf, 10, y, 190, f"{loc} ({pct:.1f}%)", val, max_val, color)
        pdf.set_y(y + 7.4)

    top5 = conteo.head(5)
    bottom5 = conteo.tail(5).sort_values()
    filas_resumen = max(len(top5), len(bottom5))
    _asegurar_espacio(pdf, 12 + filas_resumen * 6)
    pdf.ln(3)
    y2 = pdf.get_y()
    pdf.set_xy(10, y2)
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*_VERDE)
    pdf.cell(90, 6, "Top 5 localidades", ln=0)
    pdf.set_xy(110, y2)
    pdf.cell(90, 6, "Localidades con menor cobertura", ln=1)
    yy = y2 + 7
    for i in range(filas_resumen):
        if i < len(top5):
            loc, val = top5.index[i], top5.iloc[i]
            pdf.set_xy(10, yy)
            pdf.set_font("Helvetica", "B", 8)
            pdf.set_text_color(*_VERDE)
            pdf.cell(6, 5, f"{i + 1}.", ln=0)
            pdf.set_font("Helvetica", "", 8.3)
            pdf.set_text_color(*_GRIS_OSCURO)
            pdf.cell(74, 5, _pdf_safe_text(f"{loc} - {val}")[:40], ln=0)
        if i < len(bottom5):
            loc, val = bottom5.index[i], bottom5.iloc[i]
            pdf.set_xy(110, yy)
            pdf.set_font("Helvetica", "B", 8)
            pdf.set_text_color(*_ROJO)
            pdf.cell(6, 5, f"{i + 1}.", ln=0)
            pdf.set_font("Helvetica", "", 8.3)
            pdf.set_text_color(*_GRIS_OSCURO)
            pdf.cell(74, 5, _pdf_safe_text(f"{loc} - {val}")[:40], ln=0)
        yy += 6
    pdf.set_y(yy + 2)


def _pagina_alertas(pdf: FPDF, alertas: list):
    pdf.add_page()
    _titulo_pagina(pdf, "Tablero de Alertas",
                    "Alertas activas identificadas para el periodo y filtros seleccionados")

    if not alertas:
        y = pdf.get_y()
        _rect_suave(pdf, 10, y, 190, 20, color=_VERDE_EXITO_CLARO, radio=2.5)
        _insignia_icono(pdf, "check", 15, y + 5, 10, _BLANCO, _VERDE_EXITO)
        pdf.set_xy(30, y + 7.5)
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(*_VERDE_EXITO)
        pdf.cell(0, 6, "Sin alertas activas para el periodo y filtros seleccionados.", ln=1)
        pdf.set_y(y + 24)
        return

    conteo_sev = {"alta": 0, "media": 0, "baja": 0}
    for a in alertas:
        clave = a.get("severidad", "media")
        conteo_sev[clave] = conteo_sev.get(clave, 0) + 1

    y = pdf.get_y()
    resumen = [("alta", "criticas", _ROJO), ("media", "moderadas", _AMARILLO), ("baja", "bajas", _VERDE_EXITO)]
    ancho_res = 190 / 3
    for i, (clave, etiqueta, color) in enumerate(resumen):
        x = 10 + i * ancho_res
        _rect_suave(pdf, x, y, ancho_res - 4, 20, color=_tinte(color), radio=2.5)
        pdf.set_xy(x + 4, y + 3)
        pdf.set_font("Helvetica", "B", 16)
        pdf.set_text_color(*color)
        pdf.cell(30, 9, str(conteo_sev.get(clave, 0)), ln=0)
        pdf.set_xy(x + 4, y + 13)
        pdf.set_font("Helvetica", "", 7.5)
        pdf.set_text_color(*_GRIS_OSCURO)
        pdf.cell(ancho_res - 8, 4, f"Alertas {etiqueta}", ln=0)
    y += 27

    pdf.set_xy(10, y)
    pdf.set_font("Helvetica", "B", 10.5)
    pdf.set_text_color(*_VERDE)
    pdf.cell(0, 6, "Detalle de alertas", ln=1)
    y = pdf.get_y() + 1
    for a in alertas:
        _asegurar_espacio(pdf, 24)
        y = pdf.get_y()
        h = _tarjeta_alerta(pdf, 10, y, 190, a)
        pdf.set_y(y + h + 3)


def _pagina_recomendaciones(pdf: FPDF, recomendaciones_matriz: list):
    pdf.add_page()
    _titulo_pagina(pdf, "Matriz de Recomendaciones",
                    "Acciones sugeridas, impacto esperado, prioridad, responsable y horizonte de implementacion")

    y = pdf.get_y()
    pdf.set_font("Helvetica", "B", 7.2)
    pdf.set_text_color(*_GRIS_TEXTO)
    pdf.set_xy(24, y)
    pdf.cell(80, 5, "ACCION RECOMENDADA", ln=0)
    pdf.set_xy(106, y)
    pdf.cell(32, 5, "IMPACTO ESPERADO", ln=0)
    pdf.set_xy(140, y)
    pdf.cell(18, 5, "PRIORIDAD", ln=0)
    pdf.set_xy(159, y)
    pdf.cell(40, 5, "RESPONSABLE / HORIZONTE", ln=1)
    _linea(pdf, 10, pdf.get_y() + 1, 200, pdf.get_y() + 1, _GRIS_MEDIO, 0.3)
    pdf.set_y(pdf.get_y() + 3)

    if not recomendaciones_matriz:
        _texto_multilinea(pdf, 10, pdf.get_y(), 190,
                           "No se generaron recomendaciones para el periodo y filtros seleccionados.", tam=9)
        return

    for item in recomendaciones_matriz:
        _asegurar_espacio(pdf, 22)
        y = pdf.get_y()
        h = _fila_matriz_recomendacion(pdf, 10, y, 190, item)
        pdf.set_y(y + h + 3)


def _pagina_cierre(pdf: FPDF):
    pdf.add_page()
    y0 = pdf.get_y()
    _rect_suave(pdf, 10, y0, 190, 90, color=_GRIS_CLARO, radio=3)
    _ilustracion_economia_circular(pdf, cx=105, cy=y0 + 28, r=19)

    pdf.set_xy(20, y0 + 52)
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(*_VERDE)
    pdf.cell(170, 8, "Gracias por su atencion", align="C", ln=1)

    _texto_multilinea(
        pdf, 25, y0 + 64, 160,
        "Este informe fue generado automaticamente por SIGATA a partir de los datos y filtros vigentes al "
        "momento de la exportacion, en el marco del Proyecto de Inversion 8215 y del ODS 12: Produccion y "
        "Consumo Responsables.",
        tam=8.6, alto_linea=3.8, color=_GRIS_OSCURO, max_lineas=5,
    )

    y1 = y0 + 100
    pdf.set_xy(10, y1)
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(*_VERDE)
    pdf.cell(0, 6, "Canales de contacto institucional", ln=1)
    pdf.set_x(10)
    pdf.set_font("Helvetica", "", 8.6)
    pdf.set_text_color(*_GRIS_OSCURO)
    pdf.cell(0, 5.5, _pdf_safe_text(APP_ENTITY), ln=1)
    pdf.set_x(10)
    pdf.cell(0, 5.5, "Unidad Administrativa Especial de Servicios Publicos - UAESP", ln=1)


def generar_informe_ejecutivo_pdf(
    df: pd.DataFrame,
    kpi: dict,
    valor_publico: dict,
    alertas: list,
    recomendaciones: list,
    analisis_texto: str,
) -> bytes:
    """
    Genera el 'Informe Ejecutivo' en PDF con diseno tipo presentacion
    ejecutiva de firma consultora, pensado para el Comite Institucional de
    Gestion y Desempeno, Alta Direccion de la UAESP, entidades de control y
    jurados de investigacion. Incluye: portada institucional, dashboard de
    indicadores estrategicos, Valor Publico Generado (gauge), resumen
    ejecutivo tipo storytelling, graficos principales, cobertura territorial
    con ranking e intensidad de color, tablero de alertas y matriz visual de
    recomendaciones. La logica, los datos y los indicadores no cambian: solo
    cambia radicalmente su presentacion visual.

    Los parametros `recomendaciones` y `analisis_texto` se conservan para no
    alterar la firma de la funcion (usada por app.py); la version enriquecida
    de las recomendaciones y la sintesis narrativa se calculan internamente
    a partir de `df`, `kpi`, `valor_publico` y `alertas`, reutilizando el
    mismo motor de reglas de `modules.analisis` que ya produjo `recomendaciones`
    y `analisis_texto`.
    """
    pdf = InformeEjecutivoPDF()
    usuario = USUARIO_CONECTADO

    variacion = analisis.variacion_periodo(df)
    recomendaciones_matriz = analisis.generar_recomendaciones_matriz(df, kpi)
    hallazgos = analisis.sintetizar_hallazgos_ejecutivos(kpi, valor_publico, alertas, recomendaciones_matriz)

    _pagina_portada(pdf, usuario)

    pdf.set_left_margin(10)
    pdf.set_right_margin(10)
    _pagina_dashboard_kpi(pdf, kpi, variacion)
    _pagina_valor_publico(pdf, valor_publico)
    _pagina_resumen_ejecutivo(pdf, hallazgos)
    _pagina_graficos(pdf, df)
    _pagina_cobertura(pdf, df)
    _pagina_alertas(pdf, alertas)
    _pagina_recomendaciones(pdf, recomendaciones_matriz)
    _pagina_cierre(pdf)

    return bytes(pdf.output())
